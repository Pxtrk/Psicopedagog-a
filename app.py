import streamlit as st
from google import genai
import json
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# ==========================================
# LÓGICA DE GENERACIÓN DE WORD (FORMATO MEDUCA)
# ==========================================

def create_word_doc(data):
    doc = Document()
    
    # Configurar márgenes y estilo general
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Encabezado principal
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("MINISTERIO DE EDUCACIÓN\nDIRECCIÓN REGIONAL DE EDUCACIÓN\nSECUENCIA DIDÁCTICA SEMANAL / QUINCENAL")
    run.bold = True

    # Datos Generales (Usando una tabla sin bordes para alinear)
    doc.add_paragraph() # Espacio
    gen_table = doc.add_table(rows=4, cols=4)
    
    # Fila 1
    gen_table.cell(0,0).text = "ASIGNATURA:"
    gen_table.cell(0,1).text = str(data.get('asignatura', ''))
    gen_table.cell(0,2).text = "HORAS SEMANALES:"
    gen_table.cell(0,3).text = str(data.get('horas', ''))
    
    # Fila 2
    gen_table.cell(1,0).text = "DOCENTE(S):"
    gen_table.cell(1,1).text = str(data.get('docente', ''))
    gen_table.cell(1,2).text = "TRIMESTRE:"
    gen_table.cell(1,3).text = str(data.get('trimestre', ''))
    
    # Fila 3
    gen_table.cell(2,0).text = "SEMANAS:"
    semanas_texto = f"Del {data.get('sem_del', '')} al {data.get('sem_al', '')} de 2026"
    gen_table.cell(2,1).text = semanas_texto
    gen_table.cell(2,2).text = "GRADO:"
    gen_table.cell(2,3).text = str(data.get('grado', ''))
    
    # Fila 4 (Tema)
    gen_table.cell(3,0).text = "TEMA(S):"
    tema_cell = gen_table.cell(3,1)
    tema_cell.merge(gen_table.cell(3,3))
    tema_cell.text = str(data.get('tema', ''))

    # Poner en negrita las etiquetas
    for row in gen_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if ":" in para.text:
                    # Hacer negrita solo si es la etiqueta
                    text = para.text
                    para.text = ""
                    para.add_run(text).bold = True

    doc.add_paragraph() # Espacio

    # ==========================================
    # TABLA 1: Planificación (Área, Objetivos, DFA)
    # ==========================================
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Table Grid'
    
    # Fila 1: Área y Meta
    table1.cell(0,0).text = f"ÁREA(S):\n{str(data.get('area', ''))}"
    table1.cell(0,1).text = f"META DE APRENDIZAJE:\n{str(data.get('meta_aprendizaje', ''))}"
    
    # Fila 2: Objetivo (Celda combinada)
    obj_cell = table1.cell(1,0)
    obj_cell.merge(table1.cell(1,1))
    obj_cell.text = f"OBJETIVO(S) DE APRENDIZAJE:\n{str(data.get('objetivo', ''))}"
    
    # Fila 3: Indicadores y DFA
    table1.cell(2,0).text = f"INDICADORES DE LOGRO:\n{str(data.get('indicadores', ''))}"
    table1.cell(2,1).text = f"APRENDIZAJE FUNDAMENTAL (DFA):\n{str(data.get('dfa', ''))}"

    for row in table1.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if ":" in para.text:
                    parts = para.text.split(":", 1)
                    para.text = ""
                    para.add_run(parts[0] + ":").bold = True
                    para.add_run(parts[1])

    doc.add_paragraph() # Espacio

    # ==========================================
    # TABLA 2: Actividades y Evaluación
    # ==========================================
    table2 = doc.add_table(rows=3, cols=4)
    table2.style = 'Table Grid'
    
    # Fila de Encabezados Principales
    act_header = table2.cell(0,0)
    act_header.text = "ACTIVIDADES"
    act_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    act_header.paragraphs[0].runs[0].bold = True
    
    eval_header = table2.cell(0,1)
    eval_header.merge(table2.cell(0,3))
    eval_header.text = "EVALUACIÓN"
    eval_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    eval_header.paragraphs[0].runs[0].bold = True

    # Fila de Sub-encabezados
    headers_col = ["ESTRATEGIAS DE APRENDIZAJE", "EVIDENCIAS", "CRITERIOS", "TIPOS E INSTRUMENTOS"]
    for i, title in enumerate(headers_col):
        cell = table2.cell(1,i)
        cell.text = title
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True

    # Fila de Contenido
    # Col 1: Estrategias
    table2.cell(2,0).text = f"INICIO:\n{str(data.get('inicio', ''))}\n\nDESARROLLO:\n{str(data.get('desarrollo', ''))}\n\nCIERRE:\n{str(data.get('cierre', ''))}"
    
    # Col 2: Evidencias
    table2.cell(2,1).text = f"DE PRODUCTO:\n{str(data.get('evidencia_producto', ''))}\n\nDE DESEMPEÑO:\n{str(data.get('evidencia_desempeno', ''))}"
    
    # Col 3: Criterios
    table2.cell(2,2).text = f"COGNITIVOS:\n{str(data.get('criterio_cognitivo', ''))}\n\nSOCIO-AFECTIVOS:\n{str(data.get('criterio_socioafectivo', ''))}"
    
    # Col 4: Instrumentos
    tipos_eval = "Diagnóstica, Formativa, Sumativa"
    table2.cell(2,3).text = f"TIPO DE EVALUACIÓN:\n{tipos_eval}\n\nINSTRUMENTOS:\n{str(data.get('instrumentos', ''))}"

    # Aplicar negrita a las viñetas internas de la tabla 2
    etiquetas_t2 = ["INICIO:", "DESARROLLO:", "CIERRE:", "DE PRODUCTO:", "DE DESEMPEÑO:", "COGNITIVOS:", "SOCIO-AFECTIVOS:", "TIPO DE EVALUACIÓN:", "INSTRUMENTOS:"]
    for para in table2.cell(2,0).paragraphs + table2.cell(2,1).paragraphs + table2.cell(2,2).paragraphs + table2.cell(2,3).paragraphs:
        for etiqueta in etiquetas_t2:
            if para.text.startswith(etiqueta):
                texto_restante = para.text.replace(etiqueta, "", 1)
                para.text = ""
                para.add_run(etiqueta).bold = True
                para.add_run(texto_restante)

    # Guardar en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# LÓGICA DE INTELIGENCIA ARTIFICIAL (GEMINI)
# ==========================================

def generate_didactic_sequence(tema, objetivo, grado, asignatura, api_key):
    client = genai.Client(api_key=api_key)
    
    prompt = f"""
    Eres un experto en currículo educativo de Panamá (MEDUCA). 
    Crea una secuencia didáctica para {grado} en {asignatura}.
    Tema: {tema}
    Objetivo: {objetivo}
    
    Devuelve un objeto JSON válido con estas claves exactas:
    "area", "meta_aprendizaje", "dfa", "indicadores", "inicio", "desarrollo", "cierre", "evidencia_producto", "evidencia_desempeno", "criterio_cognitivo", "criterio_socioafectivo", "instrumentos"
    """
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config={
                "response_mime_type": "application/json"
            }
        )
        return json.loads(response.text)
        
    except json.JSONDecodeError as e:
        return {"error": f"Error de formato en la respuesta de la IA. Por favor, haz clic en generar de nuevo. Detalle técnico: {str(e)}"}
    except Exception as e:
        return {"error": str(e)}

# ==========================================
# INTERFAZ DE USUARIO (STREAMLIT)
# ==========================================

st.set_page_config(page_title="Planificación MEDUCA", layout="wide", page_icon="📝")

st.title("📝 Plataforma de Planificación Didáctica (Formato Oficial)")
st.markdown("Generador de secuencias estructuradas según los estándares regionales de educación.")

with st.sidebar:
    st.header("⚙️ Configuración IA")
    api_key = st.text_input("Gemini API Key", type="password")
    st.divider()
    st.header("📥 Exportar Documento")
    descarga_container = st.container()

if 'plan_data' not in st.session_state:
    st.session_state.plan_data = {}

tab_ia, tab_manual = st.tabs(["✨ Generación Automática (IA)", "✍️ Formulario Manual"])

# --- PESTAÑA: GENERACIÓN AUTOMÁTICA ---
with tab_ia:
    st.info("Ingresa los datos generales. La Inteligencia Artificial estructurará las estrategias, evidencias y criterios.")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1: docente_ia = st.text_input("Docente", placeholder="Ej. Prof. Juan Pérez")
    with col2: trimestre_ia = st.text_input("Trimestre", value="Primero")
    with col3: sem_del_ia = st.text_input("Semanas Del", placeholder="Ej. 15 de marzo")
    with col4: sem_al_ia = st.text_input("Semanas Al", placeholder="Ej. 30 de marzo")

    col_ia1, col_ia2, col_ia3 = st.columns(3)
    with col_ia1:
        asignatura_ia = st.text_input("Asignatura", value="Geografía")
    with col_ia2:
        grado_ia = st.text_input("Grado", value="9° Grado")
    with col_ia3:
        horas_ia = st.text_input("Horas Semanales", value="4")
        
    tema_ia = st.text_input("Tema a desarrollar", value="Geografía Humana y actividades económicas de Europa, Asia, Oceanía y Regiones Polares")
    objetivo_ia = st.text_area("Objetivo de Aprendizaje")
        
    if st.button("✨ Generar Secuencia Formato MEDUCA", type="primary"):
        if not api_key:
            st.error("Por favor, ingresa tu API Key de Gemini en la barra lateral.")
        elif not tema_ia or not objetivo_ia:
            st.warning("El tema y el objetivo son obligatorios.")
        else:
            with st.spinner("Diseñando actividades, DFA, criterios y evidencias..."):
                res_ia = generate_didactic_sequence(tema_ia, objetivo_ia, grado_ia, asignatura_ia, api_key)
                
                if "error" in res_ia:
                    st.error(f"Error en la generación: {res_ia['error']}")
                else:
                    st.session_state.plan_data = {
                        "docente": docente_ia, "trimestre": trimestre_ia, "sem_del": sem_del_ia, "sem_al": sem_al_ia,
                        "asignatura": asignatura_ia, "grado": grado_ia, "horas": horas_ia,
                        "tema": tema_ia, "objetivo": objetivo_ia,
                        "area": res_ia.get("area", ""),
                        "meta_aprendizaje": res_ia.get("meta_aprendizaje", ""),
                        "dfa": res_ia.get("dfa", ""),
                        "indicadores": res_ia.get("indicadores", ""),
                        "inicio": res_ia.get("inicio", ""),
                        "desarrollo": res_ia.get("desarrollo", ""),
                        "cierre": res_ia.get("cierre", ""),
                        "evidencia_producto": res_ia.get("evidencia_producto", ""),
                        "evidencia_desempeno": res_ia.get("evidencia_desempeno", ""),
                        "criterio_cognitivo": res_ia.get("criterio_cognitivo", ""),
                        "criterio_socioafectivo": res_ia.get("criterio_socioafectivo", ""),
                        "instrumentos": res_ia.get("instrumentos", "")
                    }
                    st.success("¡Secuencia generada con éxito! Descarga el archivo Word para ver la tabla oficial.")

# --- PESTAÑA: FORMULARIO MANUAL ---
with tab_manual:
    st.info("Completa los campos manualmente si no deseas usar la Inteligencia Artificial.")
    
    with st.form("form_manual"):
        st.subheader("Datos Generales")
        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
        with col_m1: docente_m = st.text_input("Docente", key="doc_m")
        with col_m2: trimestre_m = st.text_input("Trimestre", key="tri_m")
        with col_m3: sem_del_m = st.text_input("Semanas Del", key="del_m")
        with col_m4: sem_al_m = st.text_input("Semanas Al", key="al_m")

        col_m5, col_m6, col_m7 = st.columns(3)
        with col_m5: asignatura_m = st.text_input("Asignatura", key="asig_m")
        with col_m6: grado_m = st.text_input("Grado", key="gra_m")
        with col_m7: horas_m = st.text_input("Horas Semanales", key="hrs_m")
            
        tema_m = st.text_input("Tema(s)", key="tem_m")
        
        st.divider()
        st.subheader("Áreas y Objetivos")
        col_m8, col_m9 = st.columns(2)
        with col_m8:
            area_m = st.text_area("Área(s)", key="are_m")
            objetivo_m = st.text_area("Objetivo(s) de Aprendizaje", key="obj_m")
            indicadores_m = st.text_area("Indicadores de Logro", key="ind_m")
        with col_m9:
            meta_m = st.text_area("Meta de Aprendizaje", key="met_m")
            dfa_m = st.text_area("Aprendizaje Fundamental (DFA)", key="dfa_m")
            
        st.divider()
        st.subheader("Actividades y Evaluación")
        col_m10, col_m11, col_m12, col_m13 = st.columns(4)
        with col_m10:
            inicio_m = st.text_area("Inicio", key="ini_m")
            desarrollo_m = st.text_area("Desarrollo", key="des_m")
            cierre_m = st.text_area("Cierre", key="cie_m")
        with col_m11:
            ev_prod_m = st.text_area("Evidencia de Producto", key="evp_m")
            ev_des_m = st.text_area("Evidencia de Desempeño", key="evd_m")
        with col_m12:
            crit_cog_m = st.text_area("Criterios Cognitivos", key="crc_m")
            crit_soc_m = st.text_area("Criterios Socio-Afectivos", key="crs_m")
        with col_m13:
            instrumentos_m = st.text_area("Instrumentos de Evaluación", key="ins_m")

        submit_manual = st.form_submit_button("Guardar y Generar Documento", type="primary")
        
        if submit_manual:
            st.session_state.plan_data = {
                "docente": docente_m, "trimestre": trimestre_m, "sem_del": sem_del_m, "sem_al": sem_al_m,
                "asignatura": asignatura_m, "grado": grado_m, "horas": horas_m,
                "tema": tema_m, "objetivo": objetivo_m,
                "area": area_m,
                "meta_aprendizaje": meta_m,
                "dfa": dfa_m,
                "indicadores": indicadores_m,
                "inicio": inicio_m,
                "desarrollo": desarrollo_m,
                "cierre": cierre_m,
                "evidencia_producto": ev_prod_m,
                "evidencia_desempeno": ev_des_m,
                "criterio_cognitivo": crit_cog_m,
                "criterio_socioafectivo": crit_soc_m,
                "instrumentos": instrumentos_m
            }
            st.success("¡Datos guardados con éxito! Revisa el menú lateral para descargar tu documento en Word.")

# --- MANEJO DE DESCARGAS ---
if st.session_state.plan_data:
    with descarga_container:
        st.success("✅ Documento Estructurado")
        word_buf = create_word_doc(st.session_state.plan_data)
        st.download_button(
            label="📄 Descargar Plantilla Word",
            data=word_buf,
            file_name="Secuencia_Didactica_MEDUCA.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )